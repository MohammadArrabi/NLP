import os
import re
import sys
from docx import Document
import json
import matplotlib.pyplot as plt
from collections import Counter
import numpy as np


class Sentence:
    def __init__(self, sentence_text, name_speaker):
        self.sentence_text = sentence_text
        self.name_speaker = name_speaker


class Protocol:
    def __init__(self, name_protocol, number_knesset, type_protocol, number_protocol):
        self.name_protocol = name_protocol
        self.number_knesset = number_knesset
        self.type_protocol = type_protocol
        self.number_protocol = number_protocol
        self.sentences = []
        self.tokens = []


doc_word_number = {
    'אפס': 0,
    'אחד': 1,
    'אחת': 1,
    'ראשונה': 1,
    'שניים': 2,
    'שתיים': 2,
    'שנייה': 2,
    'שלוש': 3,
    'שלושה': 3,
    'ארבע': 4,
    'ארבעה': 4,
    'חמש': 5,
    'חמישה': 5,
    'שש': 6,
    'ששה': 6,
    'שבע': 7,
    'שבעה': 7,
    'שמונה': 8,
    'תשע': 9,
    'תשעה': 9,
    'עשר': 10,
    'עשרה': 10,
    'עשרים': 20,
    'שלושים': 30,
    'ארבעים': 40,
    'חמישים': 50,
    'שישים': 60,
    'שבעים': 70,
    'שמונים': 80,
    'תשעים': 90,
    'מאה': 100,
    'מאתיים': 200,
}


def extracting_data_protocol_file(file_name):
    knesset_number_number, protocol_type, file_number = file_name.split('_')
    knesset_number_number = int(knesset_number_number)
    protocol_type = "committee" if protocol_type == "ptv" else "plenary"

    return knesset_number_number, protocol_type


def ptv_protocol_number(document):
    pattern = r"פרוטוקול מס'\s+(\d+)"
    num_pattern = r"(\d+)"
    for paragraph in document.paragraphs:
        match = re.search(pattern, paragraph.text)
        if match:
            protocol_number = match.group(1)
            protocol_number = int(protocol_number)
            end_index = match.end(1)
            if (end_index < len(paragraph.text) and re.search(num_pattern, paragraph.text[end_index:])):
                txt = paragraph.text[end_index:]

            else:
                return protocol_number
            while (re.search(num_pattern, txt)):
                match1 = re.search(num_pattern, txt)
                protocol_number_con = match1.group(1)
                protocol_number_con = int(protocol_number_con)
                protocol_number = str(protocol_number) + ", " + str(protocol_number_con)
                end_index = match1.end(1)
                if (end_index < len(txt)):
                    txt = txt[end_index:]
                else:
                    break

            return protocol_number
    return -1


def word_to_number(arr_word_num):
    if (len(arr_word_num) == 1):
        return doc_word_number[arr_word_num[0]]
    else:
        calculate_num = 0
        for i in range(len(arr_word_num)):
            current_word_num = arr_word_num[i]
            if (current_word_num == "עשרה" or current_word_num == "מאות"):
                continue
            if (((i + 1) < len(arr_word_num)) and (arr_word_num[i + 1] == "עשרה" or arr_word_num[i + 1] == "מאות")):

                if (arr_word_num[i + 1] == "עשרה"):
                    calculate_num += (doc_word_number[current_word_num] + 10)
                else:
                    calculate_num += (doc_word_number[current_word_num] * 100)
            else:
                if (current_word_num[0] == 'ו'):
                    current_word_num = current_word_num[1:]
                    calculate_num += doc_word_number[current_word_num]
                else:
                    calculate_num += doc_word_number[current_word_num]

        return calculate_num


def ptm_protocol_number(document):
    pattern = r"הישיבה+\s+ה"
    for paragraph in document.paragraphs:
        match = re.search(pattern, paragraph.text)
        if match:
            end_index = paragraph.text.find(" של ")
            word_number_part = paragraph.text[8:end_index]
            arr_part_word_num = word_number_part.split("-")
            return word_to_number(arr_part_word_num)
            break
    return -1


def fetching_text_with_content(document):
    found = 0
    first_one = 0
    text = ""
    delt = 0
    own_name = ""
    pattern = r"הצבעה מס' (\d+)"
    sentences_list = []
    for paragraph in document.paragraphs:
        is_name = 0
        if (found == 0):
            if (((paragraph.text.find('היו"ר') != -1 or paragraph.text.find('היו”ר') != -1) and paragraph.text.find(
                    ":") != -1 and paragraph.text.find("מסקנות הוועדה") == -1) or paragraph.text.find(
                    "<< יור >>") != -1 or paragraph.text == 'מר שלמי גולדברג:'):
                is_name = 1
                found = 1
            elif (((paragraph.text.startswith('היו"ר') or paragraph.text.startswith(
                    'היו”ר') or paragraph.text.startswith('מ"מ היו"ר')) and (
                           paragraph.text.endswith(":") or paragraph.text.endswith(
                           ":\n"))) or paragraph.text == 'מר שלמי גולדברג:' or paragraph.text.find(
                    'יואב קיש (יו"ר הוועדה המסדרת):') != -1 or paragraph.text.find(
                    'היו"ר ניסן סלומינסקי:') != -1 or paragraph.text.find(
                    'אלי דלל (יו”ר הוועדה המיוחדת לזכויות הילד):') != -1):
                found = 1
        if found:
            if (((((paragraph.text.endswith(":") and not (
                    '0' <= paragraph.text[paragraph.text.find(":") - 1] <= '9')) or (
                           paragraph.text.find(":") != -1)) and paragraph.text != ":") and len(
                    paragraph.text) <= 35 and (paragraph.text.find("לשאול") == -1 and paragraph.text.find(
                    "להלן תשובת בנק ישראל:") == -1 and paragraph.text.find("יש לי הצעת") == -1) and paragraph.text.find(
                    "קריאה") == -1 and paragraph.text.find("לזכאות:") == -1 and paragraph.text.find(
                    "היא:") == -1 and paragraph.text.find("מסקנות הוועדה") == -1 and paragraph.text.find(
                    "הצעת חוק") == -1 and paragraph.text.find("(3)  בסופה יבוא:") == -1) or paragraph.text.find(
                    'יואב קיש (יו"ר הוועדה המסדרת):') != -1 or is_name == 1 or paragraph.text.find(
                    'אלי דלל (יו”ר הוועדה המיוחדת לזכויות הילד):') != -1 or (
                    paragraph.text.find("<< יור >>") != -1 or paragraph.text.find(
                    "<< אורח >>") != -1 or paragraph.text.find("<< דובר >>") != -1 or paragraph.text.find(
                    "<< דובר_המשך >>") != -1)):
                voting_continue = 0
                if (first_one == 1):
                    sentence = Sentence(text + "\n", own_name)
                    sentences_list.append(sentence)
                first_one = 1
                if (paragraph.text.find("אלי גולדשמידט") != -1):
                    own_name = "אלי גולדשמידט"
                    continue
                if (paragraph.text.find("אברהם פורז") != -1):
                    own_name = "אברהם פורז"
                    continue
                if (paragraph.text.find("זבולון אורלב") != -1):
                    own_name = "זבולון אורלב"
                    continue
                if (paragraph.text.find("שלום שימחון") != -1):
                    own_name = "שלום שימחון"
                    continue
                if (paragraph.text.find("גילה גמליאל") != -1):
                    own_name = "גילה גמליאל"
                    continue
                if (paragraph.text.find("מנחם בן-ששון") != -1):
                    own_name = "מנחם בן-ששון"
                    continue
                if (paragraph.text.find("ג'ודי וסרמן") != -1):
                    own_name = "ג'ודי וסרמן"
                    continue
                if (paragraph.text.find("שלי יחימוביץ'") != -1):
                    own_name = "שלי יחימוביץ'"
                    continue
                if (paragraph.text.find("עיסאווי פריג'") != -1):
                    own_name = "עיסאווי פריג'"
                    continue
                if (paragraph.text.find("בצלאל סמוטריץ'") != -1):
                    own_name = "בצלאל סמוטריץ'"
                    continue
                if (paragraph.text.find("יואב סגלוביץ'") != -1):
                    own_name = "יואב סגלוביץ'"
                    continue
                if (paragraph.text.find("מיכל לזרוביץ'") != -1):
                    own_name = "מיכל לזרוביץ'"
                    continue
                chk = paragraph.text.find("מנהל הוועדה:")
                if (chk != -1):
                    own_name = paragraph.text[chk + 13:]
                    continue
                chk = paragraph.text.find("מנהלת הוועדה")
                if (chk != -1):
                    own_name = paragraph.text[chk + 14:]
                    continue
                chk = paragraph.text.find("יועץ לוועדה:")
                if (chk != -1):
                    own_name = paragraph.text[chk + 12:]
                    continue
                if (paragraph.text.find("חברי הוועדה:") != -1):
                    chk = paragraph.text.find("-")
                    if (chk != -1):
                        own_name = paragraph.text[:chk]
                        chk -= 1
                        while (own_name[chk] == " "):
                            own_name = own_name[:chk]
                            chk -= 1
                        continue
                    else:
                        chk = paragraph.text.find("–")
                        if (chk != -1):
                            own_name = paragraph.text[:chk]
                            chk -= 1
                            while (own_name[chk] == " "):
                                own_name = own_name[:chk]
                                chk -= 1
                            continue
                end_index = paragraph.text.find("(")
                if (end_index == -1):
                    end_index = paragraph.text.find(":") + 1
                match = re.search(r" [א-ת]+' ", paragraph.text)
                match1 = re.search(r" מר ", paragraph.text)
                if match:
                    start_index = match.end()
                elif match1:
                    start_index = match1.end()
                else:
                    start_index = paragraph.text.find('היו"ר')
                    start_index1 = paragraph.text.find('היו”ר')
                    if (start_index != -1):
                        start_index = start_index + 6
                    elif (start_index1 != -1):
                        start_index = start_index1 + 6
                    else:
                        start_index = 0
                while (start_index < len(paragraph.text) and paragraph.text[start_index] == " "):
                    start_index += 1
                own_name = paragraph.text[start_index:end_index - 1]
                chk = own_name.find(">")
                if (chk != -1):
                    own_name = own_name[chk:]
                    while (own_name[0] == ">" or own_name[0] == " "):
                        own_name = own_name[1:]
                chk = own_name.find("<")
                if (chk != -1):
                    while (chk > 0 and (own_name[chk] == "<" or own_name[chk] == " ")):
                        own_name = own_name[:chk]
                        chk -= 1
                chk = own_name.find("שר הבריאות")
                if (chk != -1):
                    own_name = own_name[chk + 11:]
                chk = own_name.find("שר התחבורה")
                if (chk != -1):
                    own_name = own_name[chk + 11:]
                chk = own_name.find("שר האוצר")
                if (chk != -1):
                    own_name = own_name[chk + 9:]
                chk = own_name.find("שר המדע, התרבות והספורט")
                if (chk != -1):
                    own_name = own_name[chk + 24:]
                chk = own_name.find("שר העבודה, הרווחה והשירותים החברתיים")
                if (chk != -1):
                    own_name = own_name[chk + 37:]
                chk = own_name.find("שר לביטחון פנים")
                if (chk != -1):
                    own_name = own_name[chk + 16:]
                chk = own_name.find('עו"ד')
                if (chk != -1):
                    own_name = own_name[chk + 5:]
                chk = own_name.find("שרת החינוך")
                if (chk != -1):
                    own_name = own_name[chk + 11:]
                chk = own_name.find("שר החינוך")
                if (chk != -1):
                    own_name = own_name[chk + 10:]
                chk = own_name.find("שר הפנים")
                if (chk != -1):
                    own_name = own_name[chk + 9:]
                chk = own_name.find("שרת המשפטים")
                if (chk != -1):
                    own_name = own_name[chk + 12:]
                chk = own_name.find("שרת החוץ")
                if (chk != -1):
                    own_name = own_name[chk + 9:]
                chk = own_name.find("שר התיירות")
                if (chk != -1):
                    own_name = own_name[chk + 11:]
                chk = own_name.find("שר המשפטים")
                if (chk != -1):
                    own_name = own_name[chk + 11:]
                chk = own_name.find("שר הביטחון")
                if (chk != -1):
                    own_name = own_name[chk + 11:]
                chk = own_name.find("שר התשתיות הלאומיות")
                if (chk != -1):
                    own_name = own_name[chk + 20:]
                chk = own_name.find("שר לאיכות הסביבה")
                if (chk != -1):
                    own_name = own_name[chk + 17:]
                chk = own_name.find("ראש הממשלה")
                if (chk != -1):
                    own_name = own_name[chk + 11:]
                chk = own_name.find("שר הבינוי והשיכון")
                if (chk != -1):
                    own_name = own_name[chk + 18:]
                chk = own_name.find("שר המדע והטכנולוגיה")
                if (chk != -1):
                    own_name = own_name[chk + 20:]
                chk = own_name.find("שר החקלאות ופיתוח הכפר")
                if (chk != -1):
                    own_name = own_name[chk + 23:]
                chk = own_name.find("לקליטת העלייה")
                if (chk != -1):
                    own_name = own_name[chk + 14:]
                chk = own_name.find("שר החוץ")
                if (chk != -1):
                    own_name = own_name[chk + 8:]
                chk = own_name.find("השר")
                if (chk != -1):
                    own_name = own_name[chk + 4:]
                chk = own_name.find("מזכיר הכנסת")
                if (chk != -1):
                    own_name = own_name[chk + 12:]
                chk = own_name.find("מזכירת הכנסת")
                if (chk != -1):
                    own_name = own_name[chk + 13:]
                if (own_name[0] == "<"):
                    own_name = own_name[1:]
                if (own_name.startswith("\n")):
                    own_name = own_name[1:]
                if (own_name.endswith("\n")):
                    own_name = own_name[:len(own_name) - 1]
                text = ""
            else:
                if (paragraph.text.find("הישיבה ננעלה בשעה") == -1 and (
                        not (paragraph.text.startswith("(")) and not (paragraph.text.endswith(")"))) and (
                        len(paragraph.text) != 0 and not (paragraph.text.isspace())) and (
                        paragraph.text != "הצבעה") and delt == 0):
                    match = re.search(pattern, paragraph.text)
                    if match:
                        voting_continue = 1
                    if (voting_continue == 0):
                        if not (paragraph.alignment == 1):
                            text += paragraph.text
                if (delt == 1):
                    delt = 0
                if (paragraph.text == "הצבעה"):
                    delt = 1

    if (len(text) != 0 and not (text.isspace())):
        sentence = Sentence(text, own_name)
        sentences_list.append(sentence)

    return sentences_list


def division_into_sentences(sentences_list):
    div_sentences_list = []
    for i in range(len(sentences_list)):
        index = 0
        current_sentence = sentences_list[i]
        while (index < len(current_sentence.sentence_text)):
            end_index1 = current_sentence.sentence_text.find(".", index)
            if (end_index1 != -1 and end_index1 + 1 < len(current_sentence.sentence_text) and
                    current_sentence.sentence_text[end_index1 + 1] == '"'):
                end_index1 = current_sentence.sentence_text.find(".", end_index1 + 1)
            while (end_index1 > 0 and current_sentence.sentence_text[end_index1 - 1].isdigit() and end_index1 + 1 < len(
                    current_sentence.sentence_text) and current_sentence.sentence_text[end_index1 + 1].isdigit()):
                continue_index = end_index1 + 1
                end_index1 = current_sentence.sentence_text.find(".", continue_index)
            end_index2 = current_sentence.sentence_text.find(";", index)
            end_index3 = current_sentence.sentence_text.find("\n", index)
            end_index4 = current_sentence.sentence_text.find('"', index)
            if (end_index4 != -1 and (end_index2 == -1 or end_index3 < end_index2) and (
                    end_index1 == -1 or end_index3 < end_index1) and (end_index3 == -1 or end_index4 < end_index3)):
                startindex = current_sentence.sentence_text.find('"', end_index4 + 1)
                end_index1 = current_sentence.sentence_text.find(".", startindex)
                end_index2 = current_sentence.sentence_text.find(";", startindex)
                end_index3 = current_sentence.sentence_text.find("\n", startindex)
            if (end_index3 != -1 and (end_index2 == -1 or end_index3 < end_index2) and (
                    end_index1 == -1 or end_index3 < end_index1)):
                sen = Sentence(current_sentence.sentence_text[index:end_index3], current_sentence.name_speaker)
                div_sentences_list.append(sen)
                break
            if (end_index1 == -1 and end_index2 == -1):
                end_index = current_sentence.sentence_text.find("\n", index)
                sen = Sentence(current_sentence.sentence_text[index:end_index], current_sentence.name_speaker)
                div_sentences_list.append(sen)
                break
            elif (end_index1 == -1):
                end_index = end_index2
            elif (end_index2 == -1):
                end_index = end_index1
            else:
                end_index = min(end_index1, end_index2)
            new_sen = current_sentence.sentence_text[index:end_index + 1]
            sen = Sentence(new_sen, current_sentence.name_speaker)
            if (len(sen.sentence_text) == 0):
                break
            div_sentences_list.append(sen)
            index = end_index + 1
            if (index >= len(current_sentence.sentence_text) - 1):
                break
            while (current_sentence.sentence_text[index] == ' ' or current_sentence.sentence_text[index] == "\n"):
                index += 1
                if (index >= len(current_sentence.sentence_text)):
                    break

    return div_sentences_list


def clean_sentences(sentences_list):
    clean_sentences_list = []
    english_pattern = (r'[a-zA-Z]')
    hebrew_pattern = (r'[א-ת]')
    for sentence in sentences_list:
        found = sentence.sentence_text.find("– – –")
        if (found == -1):
            found = sentence.sentence_text.find("– –")
        if (found == -1):
            found = sentence.sentence_text.find("- - -")
        if (found == -1):
            found = sentence.sentence_text.find("- -")
        if (found == -1 and not (re.search(english_pattern, sentence.sentence_text)) and (
        re.search(hebrew_pattern, sentence.sentence_text))):
            clear_sentence = ""
            cl_sn = sentence.sentence_text
            found = sentence.sentence_text.find(r' – ')
            while (found != -1):
                clear_sentence += cl_sn[:found]
                cl_sn = cl_sn[found + 2:]
                found = cl_sn.find(r' – ')
                if (found == -1):
                    clear_sentence += cl_sn
            if (clear_sentence != ""):
                sentence.sentence_text = clear_sentence
            clean_sentences_list.append(sentence)
    return clean_sentences_list


def tokenization(sentences_list):
    tokens_list = []
    for i in range(len(sentences_list)):
        index = 0
        current_tokens_list = []
        sentence = sentences_list[i].sentence_text
        while (index < len(sentence)):
            if (index == 0 and sentence[0] == " "):
                while (sentence[index] == " "):
                    index += 1
                    if (index >= len(sentence)):
                        break
            if (index < len(sentence) and (
                    sentence[index] == '"' or sentence[index] == "'" or sentence[index] == ")" or sentence[
                index] == "(")):
                current_tokens_list.append(sentence[index])
                index += 1
            end_index = sentence.find(" ", index)
            if (end_index == -1):
                if (sentence[len(sentence) - 1] == "." or sentence[len(sentence) - 1] == "?" or sentence[
                    len(sentence) - 1] == ";" or sentence[len(sentence) - 1] == ":" or sentence[
                    len(sentence) - 1] == ")" or sentence[len(sentence) - 1] == "(" or sentence[
                    len(sentence) - 1] == '"'):
                    if (sentence[len(sentence) - 2] == '"' or sentence[len(sentence) - 2] == ')' or sentence[
                        len(sentence) - 2] == '(' or sentence[len(sentence) - 2] == "?" or sentence[
                        len(sentence) - 2] == ':'):
                        current_tokens_list.append(sentence[index:len(sentence) - 2])
                        current_tokens_list.append(sentence[len(sentence) - 2])
                        current_tokens_list.append(sentence[len(sentence) - 1])
                    else:
                        current_tokens_list.append(sentence[index:len(sentence) - 1])
                        current_tokens_list.append(sentence[len(sentence) - 1])
                    break
                else:
                    current_tokens_list.append(sentence[index:len(sentence)])
                    break
            if (sentence[end_index - 1] == "."):
                if (sentence[end_index - 2] == '"' or sentence[end_index - 2] == ')' or sentence[
                    end_index - 2] == '(' or sentence[end_index - 2] == '?' or sentence[end_index - 2] == ':'):
                    current_tokens_list.append(sentence[index:end_index - 2])
                    current_tokens_list.append(sentence[end_index - 2])
                    current_tokens_list.append(sentence[end_index - 1])
                else:
                    current_tokens_list.append(sentence[index:end_index - 1])
                    current_tokens_list.append(sentence[end_index - 1])
            elif (sentence[end_index - 1] == ";"):
                if (sentence[end_index - 2] == '"' or sentence[end_index - 2] == ')' or sentence[
                    end_index - 2] == '(' or sentence[end_index - 2] == '?' or sentence[end_index - 2] == ':'):
                    current_tokens_list.append(sentence[index:end_index - 2])
                    current_tokens_list.append(sentence[end_index - 2])
                    current_tokens_list.append(sentence[end_index - 1])
                else:
                    current_tokens_list.append(sentence[index:end_index - 1])
                    current_tokens_list.append(sentence[end_index - 1])
            elif (sentence[end_index - 1] == ","):
                if (sentence[end_index - 2] == '"' or sentence[end_index - 2] == ')' or sentence[
                    end_index - 2] == '(' or sentence[end_index - 2] == '?' or sentence[end_index - 2] == ':' or
                        sentence[end_index - 2] == "." or sentence[end_index - 2] == ";"):
                    current_tokens_list.append(sentence[index:end_index - 2])
                    current_tokens_list.append(sentence[end_index - 2])
                    current_tokens_list.append(sentence[end_index - 1])
                else:
                    current_tokens_list.append(sentence[index:end_index - 1])
                    current_tokens_list.append(sentence[end_index - 1])
            elif (sentence[end_index - 1] == "?"):
                if (sentence[end_index - 2] == '"' or sentence[end_index - 2] == ')' or sentence[
                    end_index - 2] == '(' or sentence[end_index - 2] == '.' or sentence[end_index - 2] == ':' or
                        sentence[end_index - 2] == ';'):
                    current_tokens_list.append(sentence[index:end_index - 2])
                    current_tokens_list.append(sentence[end_index - 2])
                    current_tokens_list.append(sentence[end_index - 1])
                else:
                    current_tokens_list.append(sentence[index:end_index - 1])
                    current_tokens_list.append(sentence[end_index - 1])
            elif (sentence[end_index - 1] == ":"):
                if (sentence[end_index - 2] == '"' or sentence[end_index - 2] == ')' or sentence[
                    end_index - 2] == '(' or sentence[end_index - 2] == '.' or sentence[end_index - 2] == ';' or
                        sentence[end_index - 2] == "?"):
                    current_tokens_list.append(sentence[index:end_index - 2])
                    current_tokens_list.append(sentence[end_index - 2])
                    current_tokens_list.append(sentence[end_index - 1])
                else:
                    current_tokens_list.append(sentence[index:end_index - 1])
                    current_tokens_list.append(sentence[end_index - 1])
            elif (sentence[end_index - 1] == ")"):
                if (sentence[end_index - 2] == '"' or sentence[end_index - 2] == '(' or sentence[
                    end_index - 2] == '.' or sentence[end_index - 2] == ':' or sentence[end_index - 2] == ';'):
                    current_tokens_list.append(sentence[index:end_index - 2])
                    current_tokens_list.append(sentence[end_index - 2])
                    current_tokens_list.append(sentence[end_index - 1])
                else:
                    current_tokens_list.append(sentence[index:end_index - 1])
                    current_tokens_list.append(sentence[end_index - 1])
            elif (sentence[end_index - 1] == "("):
                if (sentence[end_index - 2] == '"' or sentence[end_index - 2] == ')' or sentence[
                    end_index - 2] == '.' or sentence[end_index - 2] == ':' or sentence[end_index - 2] == ';'):
                    current_tokens_list.append(sentence[index:end_index - 2])
                    current_tokens_list.append(sentence[end_index - 2])
                    current_tokens_list.append(sentence[end_index - 1])
                else:
                    current_tokens_list.append(sentence[index:end_index - 1])
                    current_tokens_list.append(sentence[end_index - 1])
            elif (sentence[end_index - 1] == '"'):
                if (sentence[end_index - 2] == '(' or sentence[end_index - 2] == ')' or sentence[
                    end_index - 2] == '.' or sentence[end_index - 2] == ':' or sentence[end_index - 2] == ';'):
                    current_tokens_list.append(sentence[index:end_index - 2])
                    current_tokens_list.append(sentence[end_index - 2])
                    current_tokens_list.append(sentence[end_index - 1])
                else:
                    current_tokens_list.append(sentence[index:end_index - 1])
                    current_tokens_list.append(sentence[end_index - 1])
            else:
                current_tokens_list.append(sentence[index:end_index])
            index = end_index + 1
            if (index == (len(sentence) - 1)):
                current_tokens_list.append(sentence[index])
                break
            elif (index > (len(sentence) - 1)):
                break
            while (sentence[index] == " "):
                index += 1
                if (index >= len(sentence)):
                    break
        tokens_list.append(current_tokens_list)

    return tokens_list


def sentences_atleast4_tokens(tokens_list, sentences_list):
    valid_sentences = []
    valid_tokens = []
    for i in range(len(tokens_list)):
        text_sentence = ""
        if (len(tokens_list[i]) >= 4):
            valid_tokens.append(tokens_list[i])
            for token in tokens_list[i]:
                text_sentence = text_sentence + token + " "
            valid_sentences.append(Sentence(text_sentence, sentences_list[i].name_speaker))
    return valid_tokens, valid_sentences


#def process_corpus(zip_file_path):
#   corpus = []
#  with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
#     for file_name in zip_ref.namelist():
#        knesset_number, protocol_type = extracting_data_protocol_file(file_name)
#       document = Document(zip_ref.open(file_name))
#      protocol_number = ptv_protocol_number(document) if protocol_type == "committee" else ptm_protocol_number(document)
#     protocol = Protocol(file_name, knesset_number, protocol_type, protocol_number)
#    sentences_list = fetching_text_with_content(document)
#   sentences_list = division_into_sentences(sentences_list)
#  sentences_list = clean_sentences(sentences_list)
# tokens_list = tokenization(sentences_list)
#valid_tokens_list , valid_sentences_list = sentences_atleast4_tokens(tokens_list,sentences_list)
#  protocol.sentences = valid_sentences_list
# protocol.tokens = valid_tokens_list
#corpus.append(protocol)

#return corpus

def process_corpus(folder_path):
    corpus = []
    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        knesset_number, protocol_type = extracting_data_protocol_file(file_name)
        document = Document(file_path)
        protocol_number = ptv_protocol_number(document) if protocol_type == "committee" else ptm_protocol_number(
            document)
        protocol = Protocol(file_name, knesset_number, protocol_type, protocol_number)
        sentences_list = fetching_text_with_content(document)
        sentences_list = division_into_sentences(sentences_list)
        sentences_list = clean_sentences(sentences_list)
        tokens_list = tokenization(sentences_list)
        valid_tokens_list, valid_sentences_list = sentences_atleast4_tokens(tokens_list, sentences_list)
        protocol.sentences = valid_sentences_list
        protocol.tokens = valid_tokens_list
        corpus.append(protocol)

    return corpus


if __name__ == "__main__":

    if len(sys.argv) != 3:
        print('must be 2 args')
        sys.exit(1)
    folder_path = sys.argv[1]
    output_jsonl_file = sys.argv[2]
    # current_directory = os.path.dirname(os.path.abspath(__file__))
    # folder_path = os.path.join(current_directory, 'knesset_protocols')
    corpus = process_corpus(folder_path)

    # output_jsonl_file = "knesset_corpus.jsonl"
    with open(output_jsonl_file, 'w', encoding='utf-8') as file:
        for protocol in corpus:
            for sentence in protocol.sentences:
                data = {
                    "protocol_name": protocol.name_protocol,
                    "knesset_number": protocol.number_knesset,
                    "protocol_type": protocol.type_protocol,
                    "protocol_number": protocol.number_protocol,
                    "speaker_name": sentence.name_speaker,
                    "sentence_text": sentence.sentence_text
                }
                try:
                    json_str = json.dumps(data, ensure_ascii=False)
                    json_str = json_str.replace('\\"', '"')
                    file.write(json_str + '\n')
                except Exception as e:
                    print(f"Failed to write sentence from protocol {protocol.name_protocol}: {e}")
